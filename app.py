"""
MBU Study Vault - AI Powered Notes & Question Paper Repository
----------------------------------------------------------------
A beginner-friendly mini project built with Streamlit + Supabase + NLTK.

Run with:
    streamlit run app.py

Everything (UI, database, authentication, upload, NLP search) lives
in this single file on purpose, to keep the project simple.

Environment variables required (set these before running):
    SUPABASE_URL   -> your Supabase project URL
    SUPABASE_KEY   -> your Supabase service_role or anon API key

Expected Supabase setup (create these once in your Supabase project):

    Table: users
        id             bigint, primary key, identity
        full_name      text, not null
        email          text, unique, not null
        password_hash  text, not null
        created_at     text, not null

    Table: files
        id               bigint, primary key, identity
        course           text, not null
        semester         text, not null
        subject          text, not null
        category         text, not null
        display_name     text, not null
        stored_filename  text, not null   (path inside the storage bucket)
        keywords         text
        uploaded_by      text
        upload_date      text, not null

        -- OPTIONAL new columns (nullable) that power Semantic Search /
        -- RAG / Knowledge Graph / smart metadata. The app works fine
        -- without them (inserts fall back automatically), but add them
        -- for the new AI features to use real document content:
        content_text     text   (extracted text of the document)
        course_code      text
        unit             text

    Storage bucket: uploads
        A bucket named "uploads" (can be public or private - the app
        downloads files through the Supabase SDK either way).

    -- New tables added for the Faculty Portal / AI Study Assistant --
    (created once in Supabase; nothing above this line was changed)

    Table: faculty_users
        id             bigint, primary key, identity
        full_name      text, not null
        email          text, unique, not null
        password_hash  text, not null
        department     text
        created_at     text, not null

    Table: question_paper_history
        id             bigint, primary key, identity
        user_email     text, not null
        subject        text, not null
        semester       text
        marks_pattern  text
        difficulty     text
        content        text, not null
        created_at     text, not null

    Table: download_logs   (optional - only powers the Faculty
                             Dashboard's "Downloads" counter; the app
                             keeps working fine even if this table is
                             never created)
        id             bigint, primary key, identity
        file_id        bigint
        downloaded_by  text
        downloaded_at  text

    Table: search_logs   (optional - powers the new Learning Analytics
                           dashboard's "Most Searched Topics" widget;
                           search still works fine without it)
        id             bigint, primary key, identity
        query          text
        user_email     text
        created_at     text

Additional environment variable required for the AI Study Assistant /
Chat Companion / RAG features (Groq LLaMA 3.3 70B):
    GROQ_API_KEY   -> your Groq API key (https://console.groq.com)
"""

import os
import re
import io
import json
import uuid
import hashlib
from datetime import datetime

import requests
import streamlit as st
import nltk
from supabase import create_client

# ----------------------------------------------------------------------
# Extra libraries used ONLY by the "Intelligent Document Validation
# System". Each import is wrapped in a try/except so that, if a package
# or system binary is missing on a particular machine, the rest of the
# app (login, upload, search, dashboard, etc.) keeps working exactly as
# before - only the specific file type that needs the missing library
# will show a clear error message instead of crashing the whole app.
#
# Install with:
#   pip install pypdf python-docx pytesseract Pillow reportlab pdf2image
#   pip install scikit-learn networkx matplotlib
#
# For OCR (image -> text, AND scanned/handwritten PDFs -> text) you
# ALSO need the Tesseract OCR engine installed on the operating system
# itself (it is not a Python package), e.g.:
#   Windows : https://github.com/UB-Mannheim/tesseract/wiki
#   Linux   : sudo apt-get install tesseract-ocr
#   Mac     : brew install tesseract
#
# Scanned/photographed PDFs (no embedded text layer, e.g. handwritten
# notes saved as PDF) ALSO need Poppler (used by pdf2image to turn PDF
# pages into images before OCR):
#   Windows : download from https://github.com/oschwartz10612/poppler-windows/releases,
#             unzip, and add its "bin" folder to your PATH (then restart
#             the terminal so PATH changes take effect)
#   Linux   : sudo apt-get install poppler-utils
#   Mac     : brew install poppler
# ----------------------------------------------------------------------
try:
    from pypdf import PdfReader  # reading text out of PDF files
except Exception:
    try:
        from PyPDF2 import PdfReader  # fallback for older environments
    except Exception:
        PdfReader = None

try:
    import docx as python_docx  # reading text out of DOCX files
except Exception:
    python_docx = None

try:
    from reportlab.pdfgen import canvas as pdf_canvas  # building new PDFs
    from reportlab.lib.pagesizes import A4
except Exception:
    pdf_canvas = None
    A4 = None

try:
    import pytesseract  # OCR engine wrapper
    from PIL import Image  # opening/reading image files
except Exception:
    pytesseract = None
    Image = None

try:
    from pdf2image import convert_from_bytes  # rasterizes scanned PDF pages for OCR
except Exception:
    convert_from_bytes = None

# ======================================================================
# 1. PAGE CONFIG (must be the very first Streamlit command)
# ======================================================================
st.set_page_config(
    page_title="MBU Study Vault",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ======================================================================
# 2. CONSTANTS & PATHS
# ======================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

LOGO_PATH = os.path.join(BASE_DIR, "mbu_logo.png")

STORAGE_BUCKET = "uploads"

ADMIN_EMAIL = "varshaseenu2005@gmail.com"

COURSES = {
    "MCA": [f"Semester {i}" for i in range(1, 5)],
    "BCA": [f"Semester {i}" for i in range(1, 7)],
    "B.Tech": [f"Semester {i}" for i in range(1, 9)],
    "MBA": [f"Semester {i}" for i in range(1, 5)],
}

CATEGORIES = ["Handwritten Notes", "Typed Notes", "Question Paper"]

FACULTY_CATEGORIES = ["Notes", "Assignments", "Lab Manuals", "Question Papers", "Syllabus"]
MARKS_PATTERNS = ["100 Marks (5x20)", "75 Marks (5x15)", "60 Marks (Mixed)", "50 Marks (Mixed)"]
DIFFICULTY_LEVELS = ["Easy", "Medium", "Hard"]


# ======================================================================
# 3. SIMPLE NLP SETUP (Tokenization + Stopword Removal)
# ======================================================================
@st.cache_resource
def setup_nltk():
    required_packages = [
        ("tokenizers/punkt", "punkt"),
        ("tokenizers/punkt_tab", "punkt_tab"),
        ("corpora/stopwords", "stopwords"),
    ]
    for find_path, package_name in required_packages:
        try:
            nltk.data.find(find_path)
        except LookupError:
            try:
                nltk.download(package_name, quiet=True)
            except Exception:
                pass
    return True


setup_nltk()

try:
    from nltk.corpus import stopwords as nltk_stopwords
    STOPWORDS = set(nltk_stopwords.words("english"))
except Exception:
    STOPWORDS = set()

try:
    from nltk.tokenize import word_tokenize
except Exception:
    word_tokenize = None


def get_keywords(text):
    if not text:
        return []

    text = text.lower()
    text = re.sub(r"[^a-z0-9\s\+\#]", " ", text)

    if word_tokenize is not None:
        try:
            tokens = word_tokenize(text)
        except Exception:
            tokens = text.split()
    else:
        tokens = text.split()

    keywords = []
    seen = set()
    for token in tokens:
        token = token.strip()
        if token and token not in STOPWORDS and len(token) > 1 and token not in seen:
            seen.add(token)
            keywords.append(token)

    return keywords


# ======================================================================
# 3B. INTELLIGENT DOCUMENT VALIDATION SYSTEM (AI + NLP decision-making)
# ======================================================================
EDUCATIONAL_KEYWORDS = {
    "unit", "chapter", "module", "syllabus", "semester", "course",
    "subject", "topic", "lecture", "notes", "assignment", "assignments",
    "question", "questions", "answer", "answers", "exam", "examination",
    "internal", "external", "marks", "credits", "university", "college",
    "department", "dept", "faculty", "professor", "lecturer", "student",
    "students", "tutorial", "exercise", "exercises", "problem",
    "problems", "solution", "solutions", "definition", "define",
    "explain", "describe", "derive", "theorem", "proof", "formula",
    "equation", "algorithm", "program", "programming", "code", "output",
    "input", "experiment", "lab", "laboratory", "practical", "viva",
    "objective", "objectives", "outcome", "outcomes", "syllabus",
    "reference", "references", "bibliography", "textbook", "introduction",
    "conclusion", "abstract", "summary", "diagram", "figure", "table",
    "example", "examples", "concept", "concepts", "theory", "principle",
    "model", "analysis", "design", "implementation", "application",
    "advantages", "disadvantages", "types", "classification", "define",
    "compare", "difference", "list", "write", "short", "note", "marks",
    "co1", "co2", "co3", "co4", "co5", "unit-1", "unit-2", "unit-3",
    "mid", "cie", "see", "grade", "gpa", "cgpa", "curriculum", "b.tech",
    "mca", "bca", "mba", "engineering", "computer", "science", "data",
    "structures", "database", "network", "system", "systems", "software",
    "hardware", "management", "research", "paper", "thesis", "project",
}


def detect_file_type(uploaded_file):
    _, ext = os.path.splitext(uploaded_file.name.lower())
    return ext.replace(".", "").strip()


def ocr_pdf_bytes_to_text(pdf_bytes, max_pages=20):
    """
    Fallback for scanned/photographed PDFs that have no embedded text
    layer (e.g. a phone photo of handwritten notes saved as PDF).
    Rasterizes each page to an image with pdf2image, then runs the
    same Tesseract OCR already used for JPG/PNG uploads. Requires
    both the 'pdf2image' Python package AND the Poppler binaries to
    be installed on the OS (see the note near the imports at the top
    of this file). Returns "" (never raises) if either is missing so
    the rest of the app keeps working.
    """
    if convert_from_bytes is None or pytesseract is None:
        return ""
    try:
        images = convert_from_bytes(pdf_bytes)
        pages_text = []
        for image in images[:max_pages]:
            if image.mode != "RGB":
                image = image.convert("RGB")
            pages_text.append(pytesseract.image_to_string(image))
        return "\n".join(pages_text)
    except Exception:
        return ""


def extract_text_from_pdf_bytes(pdf_bytes, debug=False):
    if PdfReader is None:
        if debug:
            st.warning("Debug: pypdf/PyPDF2 is not installed - PdfReader is None.")
        return ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        pages_text = []
        for page in reader.pages:
            pages_text.append(page.extract_text() or "")
        result = "\n".join(pages_text)
        if debug:
            st.warning(
                f"Debug: pypdf read {len(reader.pages)} page(s), "
                f"extracted {len(result)} characters of embedded text."
            )
    except Exception as e:
        if debug:
            st.warning(f"Debug: pypdf raised an exception while reading this PDF: {e}")
        result = ""

    # No embedded text found (common for scanned/photographed PDFs,
    # e.g. handwritten notes) - fall back to OCR on the page images.
    if not result.strip():
        if debug and (convert_from_bytes is None or pytesseract is None):
            st.warning(
                "Debug: No embedded text found in this PDF (it looks like a "
                "scanned/photographed document). OCR fallback is unavailable "
                "because 'pdf2image' and/or Tesseract OCR aren't installed."
            )
        ocr_result = ocr_pdf_bytes_to_text(pdf_bytes)
        if debug and ocr_result.strip():
            st.warning(
                f"Debug: No embedded text found, but OCR fallback extracted "
                f"{len(ocr_result)} characters from the scanned pages."
            )
        return ocr_result

    return result


def convert_docx_to_pdf(docx_bytes):
    if python_docx is None or pdf_canvas is None:
        raise Exception(
            "DOCX support is not available on this server "
            "(missing 'python-docx' or 'reportlab' library)."
        )

    document = python_docx.Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    extracted_text = "\n".join(paragraphs)

    pdf_buffer = io.BytesIO()
    c = pdf_canvas.Canvas(pdf_buffer, pagesize=A4)
    page_width, page_height = A4
    left_margin = 50
    top_margin = page_height - 50
    line_height = 14
    max_chars_per_line = 100

    y = top_margin
    c.setFont("Helvetica", 10)
    for paragraph in paragraphs:
        lines = [paragraph[i:i + max_chars_per_line] for i in range(0, len(paragraph), max_chars_per_line)] or [""]
        for line in lines:
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = top_margin
            c.drawString(left_margin, y, line)
            y -= line_height
        y -= line_height * 0.5

    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer.read(), extracted_text


def ocr_image_to_searchable_pdf(image_bytes):
    if pytesseract is None or Image is None:
        raise Exception(
            "OCR support is not available on this server "
            "(missing 'pytesseract'/'Pillow' library or the Tesseract "
            "OCR engine)."
        )

    image = Image.open(io.BytesIO(image_bytes))
    if image.mode != "RGB":
        image = image.convert("RGB")

    extracted_text = pytesseract.image_to_string(image)
    pdf_bytes = pytesseract.image_to_pdf_or_hocr(image, extension="pdf")
    return pdf_bytes, extracted_text


def classify_document_content(text):
    keywords = get_keywords(text)

    if len(keywords) < 5:
        return False, keywords

    educational_hits = sum(1 for k in keywords if k in EDUCATIONAL_KEYWORDS)
    hit_ratio = educational_hits / len(keywords)

    is_educational = educational_hits >= 3 or hit_ratio >= 0.04

    return is_educational, keywords


def validate_and_process_document(uploaded_file):
    file_type = detect_file_type(uploaded_file)
    file_bytes = uploaded_file.getvalue()
    base_name, _ = os.path.splitext(uploaded_file.name)

    try:
        if file_type == "pdf":
            pdf_bytes = file_bytes
            extracted_text = extract_text_from_pdf_bytes(file_bytes)
            output_filename = uploaded_file.name

        elif file_type == "docx":
            pdf_bytes, extracted_text = convert_docx_to_pdf(file_bytes)
            output_filename = f"{base_name}.pdf"

        elif file_type in ("jpg", "jpeg", "png"):
            pdf_bytes, extracted_text = ocr_image_to_searchable_pdf(file_bytes)
            output_filename = f"{base_name}.pdf"

        else:
            return {
                "accepted": False,
                "message": (
                    f"Unsupported file type '.{file_type}'. Please upload a "
                    "PDF, DOCX, JPG, JPEG, or PNG file."
                ),
            }
    except Exception as e:
        return {
            "accepted": False,
            "message": f"Could not analyze this file: {e}",
        }

    is_educational, keywords = classify_document_content(extracted_text)

    if not is_educational:
        return {
            "accepted": False,
            "message": (
                "This file is not an educational document. Please upload "
                "notes, assignments, lab manuals, or question papers."
            ),
        }

    return {
        "accepted": True,
        "pdf_bytes": pdf_bytes,
        "filename": output_filename,
        "keywords": keywords,
        "content_text": extracted_text,
    }


# ======================================================================
# 4. SUPABASE CLIENT SETUP
# ======================================================================
_secrets_error = None
try:
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
except Exception as e:
    _secrets_error = e
    SUPABASE_URL = os.getenv("SUPABASE_URL")
    SUPABASE_KEY = os.getenv("SUPABASE_KEY")


@st.cache_resource
def init_supabase():
    if not SUPABASE_URL or not SUPABASE_KEY:
        st.error(
            "Missing Supabase configuration. Could not load SUPABASE_URL / "
            "SUPABASE_KEY from st.secrets, and they are not set as OS "
            "environment variables either.\n\n"
            f"Underlying error from st.secrets: `{_secrets_error}`\n\n"
            "Most common cause: Streamlit was not launched from the exact "
            "folder that contains the `.streamlit` directory."
        )
        st.stop()
    return create_client(SUPABASE_URL, SUPABASE_KEY)


supabase = init_supabase()


# ======================================================================
# 5. AUTHENTICATION HELPERS
# ======================================================================
def hash_password(password):
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def create_user(full_name, email, password):
    email = email.lower().strip()
    try:
        existing = (
            supabase.table("users")
            .select("id")
            .eq("email", email)
            .execute()
        )
        if existing.data:
            return False, "An account with this email already exists."

        supabase.table("users").insert(
            {
                "full_name": full_name.strip(),
                "email": email,
                "password_hash": hash_password(password),
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        ).execute()
        return True, "Account created successfully! Please login."
    except Exception as e:
        return False, f"Could not create account: {e}"


def verify_user(email, password):
    email = email.lower().strip()
    try:
        response = (
            supabase.table("users")
            .select("*")
            .eq("email", email)
            .execute()
        )
        if response.data:
            user = response.data[0]
            if user.get("password_hash") == hash_password(password):
                return user
    except Exception:
        pass
    return None


# ======================================================================
# 6. FILE / SEARCH HELPERS (Supabase table `files` + Storage bucket)
# ======================================================================
def save_file_record(course, semester, subject, category, display_name, stored_filename, uploaded_by, document_keywords=None, content_text=None, course_code=None, unit=None):
    keyword_source = f"{course} {semester} {subject} {category} {display_name}"
    keywords = get_keywords(keyword_source)

    if document_keywords:
        for kw in document_keywords:
            if kw not in keywords:
                keywords.append(kw)

    keywords = " ".join(keywords)

    record = {
        "course": course,
        "semester": semester,
        "subject": subject,
        "category": category,
        "display_name": display_name,
        "stored_filename": stored_filename,
        "keywords": keywords,
        "uploaded_by": uploaded_by,
        "upload_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    extra = {}
    if content_text:
        extra["content_text"] = content_text[:20000]
    if course_code:
        extra["course_code"] = course_code
    if unit:
        extra["unit"] = unit

    try:
        supabase.table("files").insert({**record, **extra}).execute()
    except Exception:
        # Optional columns (content_text/course_code/unit) probably
        # don't exist yet in this Supabase project - fall back to the
        # original, always-safe insert so uploads never break.
        supabase.table("files").insert(record).execute()


def get_subjects(course, semester):
    response = (
        supabase.table("files")
        .select("subject")
        .eq("course", course)
        .eq("semester", semester)
        .execute()
    )
    subjects = sorted({row["subject"] for row in response.data if row.get("subject")})
    return subjects


def get_files(course=None, semester=None, subject=None):
    query = supabase.table("files").select("*")
    if course:
        query = query.eq("course", course)
    if semester:
        query = query.eq("semester", semester)
    if subject:
        query = query.eq("subject", subject)
    response = query.order("upload_date", desc=True).execute()
    return response.data or []


def search_files(user_query):
    query_keywords = get_keywords(user_query)
    if not query_keywords:
        return []

    response = supabase.table("files").select("*").execute()
    rows = response.data or []

    results = []
    for row in rows:
        file_keywords = set((row.get("keywords") or "").split())
        match_count = sum(1 for k in query_keywords if k in file_keywords)

        if match_count == 0:
            combined_text = (
                f"{row['course']} {row['semester']} {row['subject']} "
                f"{row['category']} {row['display_name']}"
            ).lower()
            match_count = sum(1 for k in query_keywords if k in combined_text)

        if match_count > 0:
            row["match_score"] = match_count
            results.append(row)

    results.sort(key=lambda r: r["match_score"], reverse=True)
    return results


def upload_file_to_storage(uploaded_file):
    unique_name = f"{uuid.uuid4().hex}_{uploaded_file.name}"
    file_bytes = uploaded_file.getvalue()

    supabase.storage.from_(STORAGE_BUCKET).upload(
        unique_name,
        file_bytes,
        {"content-type": "application/pdf"},
    )
    return unique_name


def upload_bytes_to_storage(file_bytes, filename):
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    supabase.storage.from_(STORAGE_BUCKET).upload(
        unique_name,
        file_bytes,
        {"content-type": "application/pdf"},
    )
    return unique_name


def download_file_from_storage(stored_filename):
    try:
        return supabase.storage.from_(STORAGE_BUCKET).download(stored_filename)
    except Exception:
        return None


def delete_file(file_id, stored_filename):
    try:
        supabase.storage.from_(STORAGE_BUCKET).remove([stored_filename])
        supabase.table("files").delete().eq("id", file_id).execute()
        return True, "File deleted successfully."
    except Exception as e:
        return False, f"Could not delete file: {e}"


# ======================================================================
# 7. STYLING - Dark Maroon / Gold / White / Light Grey theme
# ======================================================================
def inject_css():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700&display=swap');

        html, body, [class*="css"] {
            font-family: 'Poppins', sans-serif;
        }

        .stApp {
            background-color: #F4F2EF;
        }

        section[data-testid="stSidebar"] {
            background-color: #4A0E0E;
        }
        section[data-testid="stSidebar"] * {
            color: #F5EFE0 !important;
        }
        section[data-testid="stSidebar"] .stButton > button {
            background-color: #6B1F1F;
            color: #F5D57A !important;
            border: 1px solid #D4AF37;
            border-radius: 10px;
            font-weight: 600;
            margin-bottom: 6px;
            width: 100%;
            transition: all 0.2s ease-in-out;
        }
        section[data-testid="stSidebar"] .stButton > button:hover {
            background-color: #D4AF37;
            color: #4A0E0E !important;
            transform: translateY(-2px);
        }

        .logo-placeholder {
            width: 90px; height: 90px; border-radius: 50%;
            background: linear-gradient(135deg, #D4AF37, #F5D57A);
            color: #4A0E0E; font-weight: 700; font-size: 22px;
            display: flex; align-items: center; justify-content: center;
            margin: 10px auto; border: 3px solid #F5EFE0;
        }
        .logo-placeholder-big {
            width: 130px; height: 130px; border-radius: 50%;
            background: linear-gradient(135deg, #D4AF37, #F5D57A);
            color: #4A0E0E; font-weight: 700; font-size: 32px;
            display: flex; align-items: center; justify-content: center;
            margin: 10px auto; border: 4px solid #4A0E0E;
        }

        .hero { text-align: center; padding: 6px 0 0 0; }
        .hero-title {
            color: #4A0E0E; font-size: 42px; font-weight: 700;
            margin-bottom: 0px; text-align: center;
        }
        .hero-sub {
            color: #B8860B; font-size: 18px; font-weight: 600;
            text-align: center; margin-top: 0px;
        }

        .welcome-card {
            background: #FFFFFF; border-left: 6px solid #D4AF37;
            border-radius: 14px; padding: 22px 28px; margin: 20px 0;
            box-shadow: 0 4px 14px rgba(0,0,0,0.08); color: #333333;
        }
        .welcome-card h3 { color: #4A0E0E; margin-top: 0; }

        .feature-card {
            background: #FFFFFF; border-radius: 16px; padding: 20px 14px;
            text-align: center; box-shadow: 0 4px 12px rgba(0,0,0,0.07);
            transition: transform 0.2s ease-in-out; min-height: 175px;
            border-top: 4px solid #D4AF37;
        }
        .feature-card:hover {
            transform: translateY(-6px);
            box-shadow: 0 8px 20px rgba(0,0,0,0.12);
        }
        .feature-icon { font-size: 30px; margin-bottom: 8px; }
        .feature-title { font-weight: 700; color: #4A0E0E; margin-bottom: 6px; }
        .feature-desc { font-size: 13px; color: #666666; }

        .course-card {
            background: linear-gradient(135deg, #4A0E0E, #6B1F1F);
            border-radius: 16px; padding: 26px 10px; text-align: center;
            color: #F5D57A; box-shadow: 0 4px 14px rgba(0,0,0,0.15);
            margin-bottom: 6px; font-weight: 700;
        }
        .course-icon { font-size: 26px; }
        .course-title { font-size: 18px; margin-top: 6px; }

        .subject-card {
            background: #FFFFFF; border-radius: 12px; padding: 14px;
            text-align: center; font-weight: 600; color: #4A0E0E;
            border: 1px solid #E8E0D0; margin-bottom: 4px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }

        .file-row { padding: 8px 0; }
        .file-name { font-weight: 600; color: #4A0E0E; font-size: 16px; }
        .file-meta { color: #777777; font-size: 13px; margin-top: 2px; }
        .badge {
            background: #D4AF37; color: #4A0E0E; padding: 2px 10px;
            border-radius: 20px; font-size: 11px; font-weight: 700;
        }
        .upload-date { color: #999999; font-size: 12px; padding-top: 14px; }
        .file-divider { border: none; border-top: 1px solid #E4E0D8; margin: 4px 0 14px 0; }

        div.stButton > button:first-child {
            background-color: #4A0E0E; color: #F5D57A;
            border: 1px solid #D4AF37; border-radius: 10px; font-weight: 600;
        }
        div.stButton > button:first-child:hover {
            background-color: #D4AF37; color: #4A0E0E; border: 1px solid #4A0E0E;
        }

        .stDownloadButton > button {
            background-color: #D4AF37 !important; color: #4A0E0E !important;
            border: 1px solid #4A0E0E !important; border-radius: 10px; font-weight: 700;
        }
        .stDownloadButton > button:hover {
            background-color: #4A0E0E !important; color: #F5D57A !important;
        }

        .footer {
            text-align: center; color: #999999; font-size: 12px;
            padding: 30px 0 10px 0;
        }

        h1, h2, h3 { color: #4A0E0E; }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ======================================================================
# 8. SESSION STATE DEFAULTS
# ======================================================================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user" not in st.session_state:
    st.session_state.user = None
if "page" not in st.session_state:
    st.session_state.page = "Home"
if "selected_course" not in st.session_state:
    st.session_state.selected_course = None
if "selected_semester" not in st.session_state:
    st.session_state.selected_semester = None
if "selected_subject" not in st.session_state:
    st.session_state.selected_subject = None

if "faculty_logged_in" not in st.session_state:
    st.session_state.faculty_logged_in = False
if "faculty_user" not in st.session_state:
    st.session_state.faculty_user = None

if "quiz_data" not in st.session_state:
    st.session_state.quiz_data = None
if "quiz_answers" not in st.session_state:
    st.session_state.quiz_answers = {}
if "quiz_submitted" not in st.session_state:
    st.session_state.quiz_submitted = False
if "quiz_score" not in st.session_state:
    st.session_state.quiz_score = 0

if "flashcards_data" not in st.session_state:
    st.session_state.flashcards_data = None
if "flashcard_index" not in st.session_state:
    st.session_state.flashcard_index = 0
if "flashcard_flipped" not in st.session_state:
    st.session_state.flashcard_flipped = False

# ---- New session state: AI Chat Companion ----
if "companion_chat_history" not in st.session_state:
    st.session_state.companion_chat_history = []


# ======================================================================
# 9. SIDEBAR NAVIGATION
# ======================================================================
def sidebar_nav():
    with st.sidebar:
        if os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, use_container_width=True)
        else:
            st.markdown('<div class="logo-placeholder">MBU</div>', unsafe_allow_html=True)

        st.markdown(
            "<h3 style='text-align:center;color:#D4AF37;'>Study Vault</h3>",
            unsafe_allow_html=True,
        )
        st.markdown("---")

        if st.session_state.logged_in:
            first_name = st.session_state.user["full_name"].split()[0]
            st.markdown(f"**Welcome, {first_name}!**")

            if st.button("🏠 Home", use_container_width=True):
                st.session_state.page = "Home"
                st.rerun()
            if st.button("📂 Dashboard", use_container_width=True):
                st.session_state.page = "Dashboard"
                st.rerun()
            if st.button("➕ Upload Material", use_container_width=True):
                st.session_state.page = "Upload"
                st.rerun()
            if st.button("🔍 Search", use_container_width=True):
                st.session_state.page = "Search"
                st.rerun()

            st.markdown("---")
            if st.button("🚪 Logout", use_container_width=True):
                st.session_state.logged_in = False
                st.session_state.user = None
                st.session_state.page = "Home"
                st.rerun()
        else:
            if st.button("🏠 Home", use_container_width=True):
                st.session_state.page = "Home"
                st.rerun()
            if st.button("📂 Dashboard", use_container_width=True):
                st.session_state.page = "Dashboard"
                st.rerun()
            if st.button("🔍 Search", use_container_width=True):
                st.session_state.page = "Search"
                st.rerun()

            st.markdown("---")
            if st.button("🔐 Login", use_container_width=True):
                st.session_state.page = "Login"
                st.rerun()
            if st.button("📝 Signup", use_container_width=True):
                st.session_state.page = "Signup"
                st.rerun()

        if st.session_state.logged_in or st.session_state.faculty_logged_in:
            st.markdown("---")
            if st.button("🤖 AI Study Assistant", use_container_width=True):
                st.session_state.page = "AI Study Assistant"
                st.rerun()
            if st.button("💬 AI Chat Companion", use_container_width=True):
                st.session_state.page = "AI Chat Companion"
                st.rerun()
            if st.button("📜 Question Paper History", use_container_width=True):
                st.session_state.page = "Question Paper History"
                st.rerun()

        st.markdown("---")
        st.markdown(
            "<h4 style='color:#D4AF37;margin-bottom:4px;'>Faculty Portal</h4>",
            unsafe_allow_html=True,
        )
        if st.session_state.faculty_logged_in:
            fac_first_name = st.session_state.faculty_user["full_name"].split()[0]
            st.markdown(f"**Faculty: {fac_first_name}**")
            if st.button("📊 Faculty Dashboard", use_container_width=True):
                st.session_state.page = "Faculty Dashboard"
                st.rerun()
            if st.button("⬆️ Faculty Upload", use_container_width=True):
                st.session_state.page = "Faculty Upload"
                st.rerun()
            if st.button("🚪 Faculty Logout", use_container_width=True):
                st.session_state.faculty_logged_in = False
                st.session_state.faculty_user = None
                st.session_state.page = "Home"
                st.rerun()
        else:
            if st.button("🔐 Faculty Login", use_container_width=True):
                st.session_state.page = "Faculty Login"
                st.rerun()
            if st.button("📝 Faculty Signup", use_container_width=True):
                st.session_state.page = "Faculty Signup"
                st.rerun()


# ======================================================================
# 10. SHARED UI COMPONENT - file list with download buttons
# ======================================================================
def render_file_list(files):
    if not files:
        st.info("No files found.")
        return

    is_admin = (
        st.session_state.logged_in
        and st.session_state.user
        and st.session_state.user.get("email", "").lower().strip() == ADMIN_EMAIL
    )

    for f in files:
        if is_admin:
            c1, c2, c3, c4 = st.columns([5, 2, 2, 1])
        else:
            c1, c2, c3 = st.columns([5, 2, 2])

        with c1:
            st.markdown(
                f"""
                <div class="file-row">
                    <div class="file-name">📄 {f['display_name']}</div>
                    <div class="file-meta">{f['course']} • {f['semester']} • {f['subject']} •
                        <span class="badge">{f['category']}</span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with c2:
            st.markdown(
                f"<div class='upload-date'>Uploaded: {f['upload_date'][:10]}</div>",
                unsafe_allow_html=True,
            )
        with c3:
            file_bytes = download_file_from_storage(f["stored_filename"])
            if file_bytes:
                download_name = f["display_name"]
                if not download_name.lower().endswith(".pdf"):
                    download_name += ".pdf"
                clicked = st.download_button(
                    label="⬇ Download",
                    data=file_bytes,
                    file_name=download_name,
                    mime="application/pdf",
                    key=f"download_{f['id']}",
                    use_container_width=True,
                )
                if clicked:
                    downloader = "guest"
                    if st.session_state.logged_in and st.session_state.user:
                        downloader = st.session_state.user.get("email", "guest")
                    elif st.session_state.faculty_logged_in and st.session_state.faculty_user:
                        downloader = st.session_state.faculty_user.get("email", "guest")
                    log_download(f["id"], downloader)
            else:
                st.error("File missing on server")

        if is_admin:
            with c4:
                if st.button("🗑", key=f"delete_{f['id']}", use_container_width=True):
                    success, message = delete_file(f["id"], f["stored_filename"])
                    if success:
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)

        st.markdown("<hr class='file-divider'>", unsafe_allow_html=True)


# ======================================================================
# 11. PAGES
# ======================================================================
def render_home():
    st.markdown('<div class="hero">', unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, width=140)
        else:
            st.markdown('<div class="logo-placeholder-big">MBU</div>', unsafe_allow_html=True)
        st.markdown('<h1 class="hero-title">MBU Study Vault</h1>', unsafe_allow_html=True)
        st.markdown(
            '<p class="hero-sub">AI Powered Notes &amp; Question Paper Repository</p>',
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class="welcome-card">
        <h3>Welcome to MBU Study Vault 👋</h3>
        <p>
        No more endless scrolling through WhatsApp and Telegram groups looking for notes.
        MBU Study Vault is a single place where Mohan Babu University students can
        <b>upload</b>, <b>search</b>, <b>view</b> and <b>download</b> handwritten notes,
        typed notes and previous question papers — organized by Course, Semester and
        Subject, and powered by AI semantic search and a RAG-based Study Assistant.
        </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Why MBU Study Vault?")
    features = [
        ("📚", "Organized Material", "Notes and papers sorted by Course, Semester & Subject."),
        ("🔍", "Smart Semantic Search", "Search by meaning, not just exact keywords."),
        ("🤖", "RAG AI Assistant", "Ask questions answered from your repository's own content."),
        ("⬇️", "Free Downloads", "Download any material with a single click."),
    ]
    cols = st.columns(4)
    for col, (icon, title, desc) in zip(cols, features):
        with col:
            st.markdown(
                f"""
                <div class="feature-card">
                    <div class="feature-icon">{icon}</div>
                    <div class="feature-title">{title}</div>
                    <div class="feature-desc">{desc}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)
    if not st.session_state.logged_in:
        c1, c2, c3 = st.columns([1, 1, 1])
        with c2:
            if st.button("Get Started → Login / Signup", use_container_width=True):
                st.session_state.page = "Login"
                st.rerun()


def render_signup():
    st.markdown("## 📝 Student Signup")
    with st.form("signup_form"):
        full_name = st.text_input("Full Name")
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        submitted = st.form_submit_button("Create Account")

        if submitted:
            if not full_name or not email or not password:
                st.error("Please fill all the fields.")
            elif password != confirm_password:
                st.error("Passwords do not match.")
            elif len(password) < 4:
                st.error("Password should be at least 4 characters long.")
            else:
                success, message = create_user(full_name, email, password)
                if success:
                    st.success(message)
                else:
                    st.error(message)

    st.markdown("Already have an account?")
    if st.button("Go to Login"):
        st.session_state.page = "Login"
        st.rerun()


def render_login():
    st.markdown("## 🔐 Student Login")
    with st.form("login_form"):
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")

        if submitted:
            user = verify_user(email, password)
            if user:
                st.session_state.logged_in = True
                st.session_state.user = user
                st.session_state.page = "Dashboard"
                st.success(f"Welcome back, {user['full_name']}!")
                st.rerun()
            else:
                st.error("Invalid email or password.")

    st.markdown("Don't have an account?")
    if st.button("Go to Signup"):
        st.session_state.page = "Signup"
        st.rerun()


def render_dashboard():
    st.markdown("## 📂 Dashboard")
    st.markdown("Select your **Course** to get started.")

    course_cols = st.columns(len(COURSES))
    for col, course in zip(course_cols, COURSES.keys()):
        with col:
            st.markdown(
                f"""
                <div class="course-card">
                    <div class="course-icon">🎓</div>
                    <div class="course-title">{course}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if st.button(f"Select {course}", key=f"course_{course}", use_container_width=True):
                st.session_state.selected_course = course
                st.session_state.selected_semester = None
                st.session_state.selected_subject = None
                st.rerun()

    if st.session_state.selected_course:
        st.markdown("---")
        st.markdown(f"### {st.session_state.selected_course} — Select Semester")
        semesters = COURSES[st.session_state.selected_course]
        sem_cols = st.columns(4)
        for i, sem in enumerate(semesters):
            with sem_cols[i % 4]:
                if st.button(sem, key=f"sem_{sem}", use_container_width=True):
                    st.session_state.selected_semester = sem
                    st.session_state.selected_subject = None
                    st.rerun()

    if st.session_state.selected_course and st.session_state.selected_semester:
        st.markdown("---")
        st.markdown(
            f"### {st.session_state.selected_course} — "
            f"{st.session_state.selected_semester} — Subjects"
        )
        subjects = get_subjects(st.session_state.selected_course, st.session_state.selected_semester)

        if not subjects:
            st.info("No subjects uploaded yet for this semester. Be the first to upload!")
            if st.button("➕ Upload Material for this Semester"):
                st.session_state.page = "Upload"
                st.rerun()
        else:
            subj_cols = st.columns(3)
            for i, subject in enumerate(subjects):
                with subj_cols[i % 3]:
                    st.markdown(f'<div class="subject-card">📘 {subject}</div>', unsafe_allow_html=True)
                    if st.button(f"View {subject}", key=f"subj_{subject}", use_container_width=True):
                        st.session_state.selected_subject = subject
                        st.rerun()

    if (
        st.session_state.selected_course
        and st.session_state.selected_semester
        and st.session_state.selected_subject
    ):
        st.markdown("---")
        st.markdown(f"### 📄 Files — {st.session_state.selected_subject}")
        files = get_files(
            st.session_state.selected_course,
            st.session_state.selected_semester,
            st.session_state.selected_subject,
        )
        render_file_list(files)

    render_learning_analytics()


def render_upload():
    st.markdown("## ➕ Upload Study Material")

    if not st.session_state.logged_in:
        st.warning("Please login to upload study material.")
        if st.button("Go to Login"):
            st.session_state.page = "Login"
            st.rerun()
        return

    with st.form("upload_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            course = st.selectbox("Course", list(COURSES.keys()))
        with col2:
            semester = st.selectbox("Semester", COURSES[course])

        subject = st.text_input("Subject (e.g. Python, DBMS, Cloud Computing)")
        category = st.selectbox("File Type", CATEGORIES)
        display_name = st.text_input("File Name (what should students see?)")

        col3, col4 = st.columns(2)
        with col3:
            course_code = st.text_input("Course Code (optional, e.g. CS301)", key="upload_course_code")
        with col4:
            unit = st.text_input("Unit (optional, e.g. Unit 3)", key="upload_unit")

        uploaded_file = st.file_uploader(
            "Choose a file (PDF, DOCX, JPG, JPEG or PNG)",
            type=["pdf", "docx", "jpg", "jpeg", "png"],
        )
        st.caption(
            "🤖 Every file is automatically analyzed by our AI/NLP engine "
            "before upload. DOCX files are converted to PDF, images are "
            "OCR-scanned into searchable PDFs, and only genuine "
            "educational material is accepted. Extracted text also powers "
            "Semantic Search, the RAG AI Study Assistant and the "
            "Knowledge Graph."
        )

        submitted = st.form_submit_button("Upload")

        if submitted:
            if not subject or not display_name or uploaded_file is None:
                st.error("Please fill all fields and choose a file.")
            else:
                with st.spinner("🤖 Analyzing document with AI/NLP... please wait"):
                    result = validate_and_process_document(uploaded_file)

                if not result["accepted"]:
                    st.error(result["message"])
                else:
                    try:
                        stored_filename = upload_bytes_to_storage(
                            result["pdf_bytes"], result["filename"]
                        )
                        save_file_record(
                            course,
                            semester,
                            subject.strip(),
                            category,
                            display_name.strip(),
                            stored_filename,
                            st.session_state.user["full_name"],
                            document_keywords=result["keywords"],
                            content_text=result.get("content_text"),
                            course_code=course_code.strip() if course_code else None,
                            unit=unit.strip() if unit else None,
                        )
                        st.success(
                            f"✅ Document verified as educational content! "
                            f"'{display_name}' uploaded successfully under "
                            f"{course} - {semester} - {subject}!"
                        )
                    except Exception as e:
                        st.error(f"Upload failed: {e}")


def render_search():
    st.markdown("## 🔍 Semantic Search")
    st.markdown(
        "Search by **meaning**, not just exact words. Powered by TF-IDF "
        "text similarity over every document's extracted content "
        "(falls back to keyword search automatically if needed)."
    )

    query = st.text_input(
        "Search by topic or question — e.g. \"Explain TCP congestion control\""
    )
    search_clicked = st.button("Search")

    if search_clicked or query:
        if query and query.strip():
            log_search(query.strip(), get_current_user_email())
            results = semantic_search_files(query)
            if not results:
                results = search_files(query)
            st.markdown(f"**{len(results)} result(s) found for** \"{query}\"")
            for r in results:
                score = r.get("relevance_score")
                if score is not None:
                    st.caption(
                        f"Relevance: {score:.0%} • Type: {r.get('category', 'N/A')} • "
                        f"Subject: {r.get('subject', 'N/A')}"
                    )
            render_file_list(results)
        else:
            st.info("Type a topic, keyword, or question to search.")


# ======================================================================
# 12. FACULTY PORTAL - DATA HELPERS
# ======================================================================
def create_faculty(full_name, email, password, department):
    email = email.lower().strip()
    try:
        existing = (
            supabase.table("faculty_users")
            .select("id")
            .eq("email", email)
            .execute()
        )
        if existing.data:
            return False, "A faculty account with this email already exists."

        supabase.table("faculty_users").insert(
            {
                "full_name": full_name.strip(),
                "email": email,
                "password_hash": hash_password(password),
                "department": (department or "").strip(),
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        ).execute()
        return True, "Faculty account created successfully! Please login."
    except Exception as e:
        return False, f"Could not create faculty account: {e}"


def verify_faculty(email, password):
    email = email.lower().strip()
    try:
        response = (
            supabase.table("faculty_users")
            .select("*")
            .eq("email", email)
            .execute()
        )
        if response.data:
            faculty = response.data[0]
            if faculty.get("password_hash") == hash_password(password):
                return faculty
    except Exception:
        pass
    return None


def get_faculty_uploads(faculty_name):
    try:
        response = (
            supabase.table("files")
            .select("*")
            .eq("uploaded_by", faculty_name)
            .order("upload_date", desc=True)
            .execute()
        )
        return response.data or []
    except Exception:
        return []


def log_download(file_id, downloaded_by):
    try:
        supabase.table("download_logs").insert(
            {
                "file_id": file_id,
                "downloaded_by": downloaded_by,
                "downloaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        ).execute()
    except Exception:
        pass


def get_download_count(file_ids):
    if not file_ids:
        return 0
    try:
        response = (
            supabase.table("download_logs")
            .select("id")
            .in_("file_id", file_ids)
            .execute()
        )
        return len(response.data or [])
    except Exception:
        return 0


# ======================================================================
# 13. AI STUDY ASSISTANT - LLM SETUP (Groq / LLaMA 3.3 70B)
# ======================================================================
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "openai/gpt-oss-120b"

try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except Exception:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY")


def call_ai(messages, temperature=0.4, max_tokens=2000, json_mode=False):
    if not GROQ_API_KEY:
        return None, (
            "AI features are not configured yet. Please ask the "
            "administrator to set the GROQ_API_KEY environment "
            "variable / Streamlit secret."
        )
    try:
        payload = {
            "model": GROQ_MODEL,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
        }
        if json_mode:
            payload["response_format"] = {"type": "json_object"}
        response = requests.post(
            GROQ_API_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=60,
        )
        response.raise_for_status()
        content = response.json()["choices"][0]["message"]["content"]
        return content, None
    except Exception as e:
        return None, f"AI request failed: {e}"


def parse_ai_json(raw_text):
    if not raw_text:
        return None
    cleaned = raw_text.strip()
    cleaned = re.sub(r"^```(json)?", "", cleaned).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        return None


def extract_text_from_any_upload(uploaded_file):
    file_type = detect_file_type(uploaded_file)
    file_bytes = uploaded_file.getvalue()
    try:
        if file_type == "pdf":
            return extract_text_from_pdf_bytes(file_bytes, debug=True)
        elif file_type == "docx":
            if python_docx is None:
                return ""
            document = python_docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in document.paragraphs)
        elif file_type == "txt":
            return file_bytes.decode("utf-8", errors="ignore")
        elif file_type in ("jpg", "jpeg", "png"):
            if pytesseract is None or Image is None:
                return ""
            image = Image.open(io.BytesIO(file_bytes))
            if image.mode != "RGB":
                image = image.convert("RGB")
            return pytesseract.image_to_string(image)
    except Exception:
        return ""
    return ""


AI_SYSTEM_PROMPT = (
    "You are an expert, friendly university teaching assistant helping "
    "students and faculty at an Indian engineering/management university. "
    "Be accurate, clear, and exam-focused."
)


def ai_generate_practice_questions(subject, unit, difficulty, num_questions, question_type="Mixed", bloom_level="Not specified", context_text=None):
    type_instruction = {
        "Mixed": "a mix of short answer, long answer and descriptive questions",
        "MCQ": "multiple-choice questions, each with 4 options labelled A-D and the correct answer marked",
        "Short Answer": "short-answer questions (2-3 mark style)",
        "Long Answer": "long-answer / essay-style questions (10+ mark style)",
        "Descriptive": "descriptive questions that require detailed explanation",
        "Bloom's Taxonomy": f"questions specifically targeting the '{bloom_level}' level of Bloom's Taxonomy",
        "Viva Questions": "oral viva-voce style questions a student might be asked in a lab/project viva",
    }.get(question_type, "a mix of short and long answer questions")

    prompt = (
        f"Create {num_questions} {difficulty}-level practice exam questions "
        f"for the subject '{subject}', unit/topic '{unit}'. The questions "
        f"should be {type_instruction}. Number them 1, 2, 3, ... Return "
        "only the questions in Markdown, no extra commentary."
    )
    if context_text:
        prompt += (
            "\n\nBase the questions on this reference material where "
            f"relevant:\n{context_text[:8000]}"
        )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.6,
    )


def ai_generate_question_paper(subject, semester, marks_pattern, difficulty, duration_minutes=None, question_type="Mixed"):
    duration_clause = f" Time duration: {duration_minutes} minutes." if duration_minutes else ""
    prompt = (
        f"Create a complete, well-formatted university question paper for "
        f"the subject '{subject}' ({semester}). Marks pattern: "
        f"{marks_pattern}. Difficulty: {difficulty}. Question style: "
        f"{question_type}.{duration_clause} Include a header with "
        f"subject, semester, time duration and maximum marks, organize the "
        f"paper into sections/units as appropriate, and number every "
        f"question. Return the paper in Markdown."
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=3000,
    )


def ai_summarize_notes(text, summary_type="Standard Summary"):
    style_instructions = {
        "Standard Summary": (
            "produce, using Markdown headings, exactly three sections: "
            "'### Short Summary' (2-4 sentences), '### Important Points' "
            "(bullet list) and '### Key Concepts' (bullet list of core "
            "terms with a one-line explanation each)"
        ),
        "Chapter Summary": "produce a structured chapter-level summary covering every major section, with headings",
        "Unit Summary": "produce a unit-wise summary broken into the unit's sub-topics, with headings",
        "Topic Summary": "produce a focused summary of the single most important topic in this text",
        "Revision Notes": "produce concise, exam-ready revision notes as short bullet points grouped under headings",
        "One-page Quick Notes": "produce extremely condensed one-page quick notes - only the most critical facts, formulas and definitions, as bullet points",
    }
    instruction = style_instructions.get(summary_type, style_instructions["Standard Summary"])
    prompt = (
        f"Read the following student notes and {instruction}."
        f"\n\nNOTES:\n{text[:12000]}"
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=1500,
    )


def ai_generate_important_questions(text):
    prompt = (
        "Based on the following notes, list the Top Important Questions "
        "a student should prepare for the exam. Number them by likely "
        f"importance. Return only the questions in Markdown.\n\nNOTES:\n{text[:12000]}"
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=1500,
    )


def ai_generate_quiz(topic, num_questions, difficulty):
    prompt = (
        f"Create {num_questions} {difficulty}-level multiple-choice "
        f"questions on the topic '{topic}'. Respond with ONLY a JSON "
        'object of the exact shape: {"questions": [{"question": "...", '
        '"options": ["A", "B", "C", "D"], "correct_index": 0}]}. '
        "correct_index is the 0-based index into options. No extra text."
    )
    content, error = call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=2500,
        json_mode=True,
    )
    if error:
        return None, error
    parsed = parse_ai_json(content)
    if not parsed or "questions" not in parsed:
        return None, "The AI response could not be understood. Please try again."
    return parsed["questions"], None


def ai_generate_flashcards(topic, num_cards):
    prompt = (
        f"Create {num_cards} revision flash cards on the topic '{topic}'. "
        'Respond with ONLY a JSON object of the exact shape: '
        '{"flashcards": [{"question": "...", "answer": "..."}]}. '
        "Keep answers short (1-2 sentences). No extra text."
    )
    content, error = call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=2000,
        json_mode=True,
    )
    if error:
        return None, error
    parsed = parse_ai_json(content)
    if not parsed or "flashcards" not in parsed:
        return None, "The AI response could not be understood. Please try again."
    return parsed["flashcards"], None


def ai_explain_topic(topic):
    prompt = (
        f"Explain the topic '{topic}' in simple, beginner-friendly English, "
        "as if teaching a first-year student. Use short paragraphs and, "
        "where useful, a small example. Return Markdown."
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=1200,
    )


def ai_ask_from_notes(notes_text, question):
    prompt = (
        "You must answer the QUESTION using ONLY the information present "
        "in the NOTES below. If the answer is not present in the notes, "
        "reply exactly: 'This isn't covered in the uploaded notes.' Do not "
        f"use outside knowledge.\n\nNOTES:\n{notes_text[:12000]}\n\n"
        f"QUESTION: {question}"
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=800,
    )


def ai_rag_answer(question, retrieved_chunks):
    """
    Retrieval-Augmented Generation: answers `question` using ONLY the
    text in `retrieved_chunks` (list of {source, text, score}). If the
    retrieved material doesn't contain the answer, the model is
    instructed to say so plainly instead of inventing one. Callers can
    show `retrieved_chunks` alongside the answer as Explainable-AI
    supporting sources.
    """
    if not retrieved_chunks:
        return (
            "No relevant material was found in the repository for this "
            "question. Try rephrasing, or upload the relevant notes first."
        ), None

    context_blocks = []
    for i, chunk in enumerate(retrieved_chunks, start=1):
        context_blocks.append(f"[Source {i}: {chunk['source']}]\n{chunk['text'][:1500]}")
    context = "\n\n".join(context_blocks)

    prompt = (
        "Answer the QUESTION using ONLY the information in the CONTEXT "
        "below, which was retrieved from documents in the repository. "
        "Cite which Source number(s) support each part of your answer "
        "using bracket notation like [Source 1]. If the context does not "
        "contain the answer, reply exactly: 'This information was not "
        "found in the uploaded repository.' Do not use outside "
        f"knowledge.\n\nCONTEXT:\n{context}\n\nQUESTION: {question}"
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=1000,
    )


def ai_generate_study_planner(subject, exam_date, hours_per_day):
    today_str = datetime.now().strftime("%Y-%m-%d")
    prompt = (
        f"Today's date is {today_str}. The exam for '{subject}' is on "
        f"{exam_date}. The student has {hours_per_day} study hours "
        "available per day. Create a day-by-day study timetable from "
        "today until the exam date, breaking the subject into a logical "
        'sequence of topics. Respond with ONLY a JSON object of the exact '
        'shape: {"plan": [{"date": "YYYY-MM-DD", "day_number": 1, '
        '"hours": 2, "topics": "..."}]}. No extra text.'
    )
    content, error = call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=2500,
        json_mode=True,
    )
    if error:
        return None, error
    parsed = parse_ai_json(content)
    if not parsed or "plan" not in parsed:
        return None, "The AI response could not be understood. Please try again."
    return parsed["plan"], None


def ai_extract_keywords(text):
    prompt = (
        "Analyze the following notes and extract: important keywords, "
        "key definitions, and technical terms. Respond with ONLY a JSON "
        'object of the exact shape: {"keywords": ["..."], "definitions": '
        '[{"term": "...", "definition": "..."}], "technical_terms": '
        f'["..."]}}. No extra text.\n\nNOTES:\n{text[:12000]}'
    )
    content, error = call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=1500,
        json_mode=True,
    )
    if error:
        return None, error
    parsed = parse_ai_json(content)
    if not parsed:
        return None, "The AI response could not be understood. Please try again."
    return parsed, None


def ai_generate_knowledge_graph_data(topic):
    prompt = (
        f"Break down the academic topic '{topic}' into a concept "
        "hierarchy/knowledge graph, the way it would be taught in a "
        "university course. Include the main topic, its sub-topics, and "
        "important related concepts (roughly 8-18 nodes total). Respond "
        'with ONLY a JSON object of the exact shape: {"nodes": ["..."], '
        '"edges": [{"source": "...", "target": "...", "label": "..."}]}. '
        "Every source/target must exactly match a string in 'nodes'. "
        "'label' should be a short relationship word like 'includes', "
        "'depends on', 'type of'. No extra text."
    )
    content, error = call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=1500,
        json_mode=True,
    )
    if error:
        return None, error
    parsed = parse_ai_json(content)
    if not parsed or "nodes" not in parsed or "edges" not in parsed:
        return None, "The AI response could not be understood. Please try again."
    return parsed, None


def ai_recommend_search_directions(topic, repo_matches):
    """
    Suggests related sub-topics and honest search DIRECTIONS (not
    fabricated book titles/authors/links, which the model can't
    verify) based on a topic and what's already in the repository.
    """
    repo_titles = ", ".join(m["display_name"] for m in repo_matches[:5]) or "none found"
    prompt = (
        f"A student is studying the topic '{topic}'. Repository resources "
        f"already found: {repo_titles}. Suggest: (1) 3-5 closely related "
        "sub-topics they should also study, (2) 3-5 good search queries "
        "they could type into Google Scholar or YouTube to find further "
        "reading/lecture videos on this topic (do NOT invent specific "
        "book titles, authors, or URLs - only suggest search phrases and "
        "general subject areas), and (3) what kind of official reference "
        "(e.g. 'your course textbook chapter on X', 'IEEE papers on Y') "
        "would help most. Return concise Markdown with headings."
    )
    return call_ai(
        [
            {"role": "system", "content": AI_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=800,
    )


COMPANION_SYSTEM_PROMPT = (
    "You are the MBU Study Vault AI Study Companion for university "
    "students. You provide three kinds of support:\n"
    "1. EDUCATIONAL SUPPORT: explain concepts simply, answer academic "
    "and programming doubts, help build study plans and revision "
    "schedules, identify important topics, and give step-by-step "
    "explanations.\n"
    "2. MOTIVATIONAL & EMOTIONAL SUPPORT: when a student expresses "
    "stress, disappointment, low motivation, or fear about exams/"
    "results, respond warmly and without judgment, acknowledge how "
    "they feel, and gently help them turn the setback into a small, "
    "concrete next study step. Do not make unrealistic promises, and "
    "do not present yourself as a therapist or medical professional. "
    "If a student's message suggests a serious safety concern (for "
    "example self-harm or crisis), respond with care, take it "
    "seriously, and clearly encourage them to reach out to a trusted "
    "person, university counsellor, or a crisis helpline right away, "
    "rather than trying to resolve it yourself.\n"
    "3. PERSONALIZED STUDY COMPANION: use the ongoing conversation to "
    "remember the student's subject, topic, and goals within this "
    "session, and tailor suggestions accordingly.\n"
    "Keep replies focused, encouraging, and exam-relevant."
)


def ai_companion_reply(chat_history, rag_context=None):
    messages = [{"role": "system", "content": COMPANION_SYSTEM_PROMPT}]
    if rag_context:
        messages.append({
            "role": "system",
            "content": (
                "The student's most recent question should be answered "
                "using this retrieved material from their uploaded "
                f"documents where relevant:\n\n{rag_context}"
            ),
        })
    messages.extend(chat_history[-12:])
    return call_ai(messages, temperature=0.5, max_tokens=900)


# ======================================================================
# 13N. SEMANTIC SEARCH + RAG RETRIEVAL (TF-IDF based)
#
#   True dense-embedding search would need a vector database and an
#   embedding model; to keep this app dependency-light and free to run
#   (no extra API keys/services), semantic similarity is computed with
#   scikit-learn's TF-IDF + cosine similarity over each document's
#   extracted content_text (falls back to keywords/metadata if
#   content_text isn't available for older uploads). This goes well
#   beyond simple exact-keyword matching and is used both for the
#   Search page and for RAG retrieval in the AI Study Assistant / Chat
#   Companion. If scikit-learn is not installed, every function below
#   degrades gracefully to the existing keyword-based search_files().
# ======================================================================
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False


@st.cache_data(ttl=60)
def _get_all_files_for_search():
    try:
        response = supabase.table("files").select("*").execute()
        return response.data or []
    except Exception:
        return []


def _file_corpus_text(row):
    parts = [
        row.get("content_text") or "",
        row.get("keywords") or "",
        row.get("display_name") or "",
        row.get("subject") or "",
        row.get("course") or "",
        row.get("category") or "",
        row.get("unit") or "",
        row.get("course_code") or "",
    ]
    return " ".join(p for p in parts if p)


def semantic_search_files(user_query, top_k=20):
    """
    TF-IDF + cosine-similarity semantic-style search across every
    uploaded file's extracted content. Returns files annotated with
    'relevance_score' (0-1), sorted by relevance. Returns [] (letting
    the caller fall back to search_files()) if scikit-learn isn't
    available or there isn't enough data to build a meaningful model.
    """
    if not SKLEARN_AVAILABLE or not user_query or not user_query.strip():
        return []

    rows = _get_all_files_for_search()
    if not rows:
        return []

    corpus = [_file_corpus_text(r) for r in rows]
    if not any(c.strip() for c in corpus):
        return []

    try:
        vectorizer = TfidfVectorizer(stop_words="english", max_features=5000)
        doc_vectors = vectorizer.fit_transform(corpus + [user_query])
        query_vector = doc_vectors[-1]
        doc_vectors = doc_vectors[:-1]
        similarities = cosine_similarity(query_vector, doc_vectors)[0]
    except Exception:
        return []

    scored = []
    for row, score in zip(rows, similarities):
        if score > 0.03:
            row = dict(row)
            row["relevance_score"] = float(score)
            scored.append(row)

    scored.sort(key=lambda r: r["relevance_score"], reverse=True)
    return scored[:top_k]


def build_rag_context(query, top_k=5):
    """
    Retrieval step for RAG: finds the top_k most relevant uploaded
    documents for `query` and returns a list of
    {"source": display_name, "text": snippet, "score": float} dicts
    ready to be handed to ai_rag_answer() / shown as Explainable-AI
    sources. Falls back to keyword search_files() if semantic search
    is unavailable.
    """
    results = semantic_search_files(query, top_k=top_k)
    if not results:
        results = search_files(query)[:top_k]
        for r in results:
            r["relevance_score"] = None

    chunks = []
    for r in results:
        text = (r.get("content_text") or r.get("keywords") or "")
        if not text.strip():
            continue
        chunks.append({
            "source": f"{r.get('display_name', 'Untitled')} ({r.get('subject', '')})",
            "text": text,
            "score": r.get("relevance_score"),
        })
    return chunks


# ======================================================================
# 13O. LEARNING ANALYTICS - search logging
# ======================================================================
def log_search(query, user_email):
    try:
        supabase.table("search_logs").insert(
            {
                "query": query,
                "user_email": user_email or "guest",
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        ).execute()
    except Exception:
        pass


def get_top_searches(limit=10):
    try:
        response = supabase.table("search_logs").select("query").execute()
        rows = response.data or []
    except Exception:
        return []
    counts = {}
    for r in rows:
        q = (r.get("query") or "").strip().lower()
        if q:
            counts[q] = counts.get(q, 0) + 1
    return sorted(counts.items(), key=lambda x: x[1], reverse=True)[:limit]


def render_learning_analytics():
    """Upgrade to the existing Dashboard: adds a Learning Analytics
    section (most searched topics, subject-wise upload distribution,
    resource type distribution). Degrades gracefully if optional
    tables (search_logs, download_logs) don't exist yet."""
    if not (st.session_state.logged_in or st.session_state.faculty_logged_in):
        return

    with st.expander("📊 Learning Analytics", expanded=False):
        all_files = _get_all_files_for_search()

        st.markdown("**Subject-wise Material Count**")
        subject_counts = {}
        for f in all_files:
            subj = f.get("subject") or "Unknown"
            subject_counts[subj] = subject_counts.get(subj, 0) + 1
        if subject_counts:
            st.bar_chart(subject_counts)
        else:
            st.caption("No uploads yet.")

        st.markdown("**Most Searched Topics**")
        top_searches = get_top_searches()
        if top_searches:
            st.bar_chart({q: c for q, c in top_searches})
        else:
            st.caption(
                "No search history yet (or the optional `search_logs` "
                "table hasn't been created in Supabase)."
            )

        st.markdown("**Resource Type Distribution**")
        category_counts = {}
        for f in all_files:
            cat = f.get("category") or "Unknown"
            category_counts[cat] = category_counts.get(cat, 0) + 1
        if category_counts:
            st.bar_chart(category_counts)


# ======================================================================
# 13P. KNOWLEDGE GRAPH RENDERING
# ======================================================================
try:
    import networkx as nx
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    GRAPH_LIBS_AVAILABLE = True
except Exception:
    GRAPH_LIBS_AVAILABLE = False


def render_knowledge_graph(graph_data):
    """Renders a {"nodes": [...], "edges": [...]} dict as a concept
    graph using networkx + matplotlib. Falls back to a readable text
    hierarchy if those libraries aren't installed."""
    nodes = graph_data.get("nodes", [])
    edges = graph_data.get("edges", [])

    if not GRAPH_LIBS_AVAILABLE:
        st.info(
            "Install `networkx` and `matplotlib` for a visual graph. "
            "Showing a text outline instead:"
        )
        for e in edges:
            st.markdown(f"- **{e.get('source')}** —({e.get('label', '')})→ **{e.get('target')}**")
        return

    try:
        G = nx.DiGraph()
        for n in nodes:
            G.add_node(n)
        for e in edges:
            if e.get("source") in nodes and e.get("target") in nodes:
                G.add_edge(e["source"], e["target"], label=e.get("label", ""))

        fig, ax = plt.subplots(figsize=(10, 7))
        pos = nx.spring_layout(G, k=1.1, seed=42)
        nx.draw_networkx_nodes(G, pos, node_color="#D4AF37", node_size=1800, ax=ax)
        nx.draw_networkx_labels(G, pos, font_size=8, font_weight="bold", font_color="#4A0E0E", ax=ax)
        nx.draw_networkx_edges(G, pos, edge_color="#6B1F1F", arrows=True, arrowsize=15, ax=ax)
        edge_labels = nx.get_edge_attributes(G, "label")
        nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=7, ax=ax)
        ax.axis("off")
        st.pyplot(fig)
        plt.close(fig)
    except Exception as e:
        st.error(f"Could not render the graph visually: {e}")
        for e2 in edges:
            st.markdown(f"- **{e2.get('source')}** —({e2.get('label', '')})→ **{e2.get('target')}**")


# ======================================================================
# 14. QUESTION PAPER HISTORY - DATA HELPERS
# ======================================================================
def save_question_paper_history(user_email, subject, semester, marks_pattern, difficulty, content):
    try:
        supabase.table("question_paper_history").insert(
            {
                "user_email": user_email,
                "subject": subject,
                "semester": semester,
                "marks_pattern": marks_pattern,
                "difficulty": difficulty,
                "content": content,
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
        ).execute()
    except Exception:
        pass


def get_question_paper_history(user_email):
    try:
        response = (
            supabase.table("question_paper_history")
            .select("*")
            .eq("user_email", user_email)
            .order("created_at", desc=True)
            .execute()
        )
        return response.data or []
    except Exception:
        return []


def delete_question_paper_history(entry_id):
    try:
        supabase.table("question_paper_history").delete().eq("id", entry_id).execute()
        return True
    except Exception:
        return False


def get_current_user_email():
    if st.session_state.logged_in and st.session_state.user:
        return st.session_state.user.get("email", "")
    if st.session_state.faculty_logged_in and st.session_state.faculty_user:
        return st.session_state.faculty_user.get("email", "")
    return ""


# ======================================================================
# 15. NEW PAGES - FACULTY PORTAL
# ======================================================================
def render_faculty_signup():
    st.markdown("## 📝 Faculty Signup")
    with st.form("faculty_signup_form"):
        full_name = st.text_input("Full Name", key="fac_signup_name")
        email = st.text_input("Email", key="fac_signup_email")
        department = st.text_input("Department", key="fac_signup_dept")
        password = st.text_input("Password", type="password", key="fac_signup_pw")
        confirm_password = st.text_input("Confirm Password", type="password", key="fac_signup_cpw")
        submitted = st.form_submit_button("Create Faculty Account")

        if submitted:
            if not full_name or not email or not password:
                st.error("Please fill all the fields.")
            elif password != confirm_password:
                st.error("Passwords do not match.")
            elif len(password) < 4:
                st.error("Password should be at least 4 characters long.")
            else:
                success, message = create_faculty(full_name, email, password, department)
                if success:
                    st.success(message)
                else:
                    st.error(message)

    st.markdown("Already have a faculty account?")
    if st.button("Go to Faculty Login"):
        st.session_state.page = "Faculty Login"
        st.rerun()


def render_faculty_login():
    st.markdown("## 🔐 Faculty Login")
    with st.form("faculty_login_form"):
        email = st.text_input("Email", key="fac_login_email")
        password = st.text_input("Password", type="password", key="fac_login_pw")
        submitted = st.form_submit_button("Login")

        if submitted:
            faculty = verify_faculty(email, password)
            if faculty:
                st.session_state.faculty_logged_in = True
                st.session_state.faculty_user = faculty
                st.session_state.page = "Faculty Dashboard"
                st.success(f"Welcome, {faculty['full_name']}!")
                st.rerun()
            else:
                st.error("Invalid email or password.")

    st.markdown("Don't have a faculty account?")
    if st.button("Go to Faculty Signup"):
        st.session_state.page = "Faculty Signup"
        st.rerun()


def render_faculty_upload():
    st.markdown("## ⬆️ Faculty Upload")

    if not st.session_state.faculty_logged_in:
        st.warning("Please login as faculty to upload material.")
        if st.button("Go to Faculty Login"):
            st.session_state.page = "Faculty Login"
            st.rerun()
        return

    with st.form("faculty_upload_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            course = st.selectbox("Course", list(COURSES.keys()), key="fac_course")
        with col2:
            semester = st.selectbox("Semester", COURSES[course], key="fac_sem")

        subject = st.text_input("Subject (e.g. Python, DBMS, Cloud Computing)", key="fac_subject")
        category = st.selectbox("Material Type", FACULTY_CATEGORIES, key="fac_category")
        display_name = st.text_input("File Name (what should students see?)", key="fac_display_name")

        col3, col4 = st.columns(2)
        with col3:
            course_code = st.text_input("Course Code (optional, e.g. CS301)", key="fac_course_code")
        with col4:
            unit = st.text_input("Unit (optional, e.g. Unit 3)", key="fac_unit")

        uploaded_file = st.file_uploader(
            "Choose a file (PDF, DOCX, JPG, JPEG or PNG)",
            type=["pdf", "docx", "jpg", "jpeg", "png"],
            key="fac_file",
        )
        st.caption(
            "🤖 Just like the student upload flow, every file is analyzed "
            "by the same AI/NLP engine before it is saved, and organized "
            "under Course → Semester → Subject automatically."
        )

        submitted = st.form_submit_button("Upload")

        if submitted:
            if not subject or not display_name or uploaded_file is None:
                st.error("Please fill all fields and choose a file.")
            else:
                with st.spinner("🤖 Analyzing document with AI/NLP... please wait"):
                    result = validate_and_process_document(uploaded_file)

                if not result["accepted"]:
                    st.error(result["message"])
                else:
                    try:
                        stored_filename = upload_bytes_to_storage(
                            result["pdf_bytes"], result["filename"]
                        )
                        save_file_record(
                            course,
                            semester,
                            subject.strip(),
                            category,
                            display_name.strip(),
                            stored_filename,
                            st.session_state.faculty_user["full_name"],
                            document_keywords=result["keywords"],
                            content_text=result.get("content_text"),
                            course_code=course_code.strip() if course_code else None,
                            unit=unit.strip() if unit else None,
                        )
                        st.success(
                            f"✅ '{display_name}' uploaded successfully under "
                            f"{course} - {semester} - {subject}!"
                        )
                    except Exception as e:
                        st.error(f"Upload failed: {e}")


def render_faculty_dashboard():
    st.markdown("## 📊 Faculty Dashboard")

    if not st.session_state.faculty_logged_in:
        st.warning("Please login as faculty to view your dashboard.")
        if st.button("Go to Faculty Login"):
            st.session_state.page = "Faculty Login"
            st.rerun()
        return

    faculty_name = st.session_state.faculty_user["full_name"]
    uploads = get_faculty_uploads(faculty_name)

    total = len(uploads)
    notes_count = sum(1 for u in uploads if u.get("category") == "Notes")
    assignments_count = sum(1 for u in uploads if u.get("category") == "Assignments")
    lab_count = sum(1 for u in uploads if u.get("category") == "Lab Manuals")
    qp_count = sum(1 for u in uploads if u.get("category") in ("Question Papers", "Question Paper"))
    syllabus_count = sum(1 for u in uploads if u.get("category") == "Syllabus")
    download_count = get_download_count([u["id"] for u in uploads])

    c1, c2, c3, c4, c5, c6, c7 = st.columns(7)
    c1.metric("Total Uploads", total)
    c2.metric("Notes", notes_count)
    c3.metric("Assignments", assignments_count)
    c4.metric("Lab Manuals", lab_count)
    c5.metric("Question Papers", qp_count)
    c6.metric("Syllabus", syllabus_count)
    c7.metric("Downloads", download_count)

    st.markdown("---")
    st.markdown("### 🕓 Recently Uploaded Files")
    render_file_list(uploads[:10])

    render_learning_analytics()


# ======================================================================
# 16. NEW PAGE - QUESTION PAPER HISTORY
# ======================================================================
def render_question_paper_history():
    st.markdown("## 📜 Question Paper History")

    if not (st.session_state.logged_in or st.session_state.faculty_logged_in):
        st.warning("Please login to view your question paper history.")
        return

    email = get_current_user_email()
    entries = get_question_paper_history(email)

    if not entries:
        st.info(
            "No AI-generated question papers yet. Create one from "
            "AI Study Assistant → AI Question Paper Generator."
        )
        return

    for entry in entries:
        with st.expander(
            f"📄 {entry['subject']} — {entry.get('semester', '')} — "
            f"{entry['created_at'][:10]}"
        ):
            st.markdown(entry["content"])
            c1, c2 = st.columns(2)
            with c1:
                st.download_button(
                    "⬇ Download as TXT",
                    data=entry["content"],
                    file_name=f"{entry['subject']}_question_paper.txt",
                    mime="text/plain",
                    key=f"qp_download_{entry['id']}",
                    use_container_width=True,
                )
            with c2:
                if st.button("🗑 Delete", key=f"qp_delete_{entry['id']}", use_container_width=True):
                    if delete_question_paper_history(entry["id"]):
                        st.success("Deleted.")
                        st.rerun()


# ======================================================================
# 17. NEW PAGE - AI STUDY ASSISTANT (available to every logged-in user)
# ======================================================================
QUESTION_TYPES = ["Mixed", "MCQ", "Short Answer", "Long Answer", "Descriptive", "Bloom's Taxonomy", "Viva Questions"]
BLOOM_LEVELS = ["Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"]
SUMMARY_TYPES = ["Standard Summary", "Chapter Summary", "Unit Summary", "Topic Summary", "Revision Notes", "One-page Quick Notes"]


def render_ai_question_generator_tool():
    """AI Study Assistant > AI Question Generator (upgraded: question
    types, Bloom's Taxonomy level, and optional grounding in an
    uploaded reference document)."""
    st.markdown("### ❓ AI Question Generator")
    with st.form("ai_qgen_form"):
        subject = st.text_input("Subject", key="qgen_subject")
        unit = st.text_input("Unit / Topic", key="qgen_unit")
        col1, col2 = st.columns(2)
        with col1:
            difficulty = st.selectbox("Difficulty", DIFFICULTY_LEVELS, key="qgen_difficulty")
        with col2:
            question_type = st.selectbox("Question Type", QUESTION_TYPES, key="qgen_type")
        bloom_level = "Not specified"
        if question_type == "Bloom's Taxonomy":
            bloom_level = st.selectbox("Bloom's Taxonomy Level", BLOOM_LEVELS, key="qgen_bloom")
        num_questions = st.number_input("Number of Questions", 1, 20, 5, key="qgen_num")
        reference_file = st.file_uploader(
            "Optional: ground questions in an uploaded document",
            type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
            key="qgen_reference",
        )
        submitted = st.form_submit_button("Generate Questions")

    if submitted:
        if not subject or not unit:
            st.error("Please fill in Subject and Unit/Topic.")
        else:
            context_text = None
            if reference_file is not None:
                context_text = extract_text_from_any_upload(reference_file)
            with st.spinner("🤖 Generating questions..."):
                content, error = ai_generate_practice_questions(
                    subject, unit, difficulty, int(num_questions),
                    question_type=question_type, bloom_level=bloom_level,
                    context_text=context_text,
                )
            if error:
                st.error(error)
            else:
                st.markdown(content)


def render_ai_question_paper_generator_tool():
    """AI Study Assistant > AI Question Paper Generator (upgraded with
    question style and optional exam duration)."""
    st.markdown("### 📄 AI Question Paper Generator")
    with st.form("ai_qpgen_form"):
        subject = st.text_input("Subject", key="qpgen_subject")
        semester = st.selectbox(
            "Semester", [f"Semester {i}" for i in range(1, 9)], key="qpgen_semester"
        )
        marks_pattern = st.selectbox("Marks Pattern", MARKS_PATTERNS, key="qpgen_marks")
        col1, col2 = st.columns(2)
        with col1:
            difficulty = st.selectbox("Difficulty", DIFFICULTY_LEVELS, key="qpgen_difficulty")
        with col2:
            question_type = st.selectbox("Question Style", QUESTION_TYPES, key="qpgen_type")
        duration_minutes = st.number_input("Exam Duration (minutes, optional)", 0, 240, 180, key="qpgen_duration")
        submitted = st.form_submit_button("Generate Question Paper")

    if submitted:
        if not subject:
            st.error("Please enter a subject.")
        else:
            with st.spinner("🤖 Generating question paper..."):
                content, error = ai_generate_question_paper(
                    subject, semester, marks_pattern, difficulty,
                    duration_minutes=duration_minutes if duration_minutes else None,
                    question_type=question_type,
                )
            if error:
                st.error(error)
            else:
                st.markdown(content)
                save_question_paper_history(
                    get_current_user_email(), subject, semester, marks_pattern, difficulty, content
                )
                st.success("✅ Saved to your Question Paper History.")


def render_ai_notes_summarizer_tool():
    """AI Study Assistant > AI Notes Summarizer (upgraded with summary
    type selector: Chapter / Unit / Topic / Revision Notes / Quick
    Notes)."""
    st.markdown("### 📝 AI Notes Summarizer")
    uploaded_file = st.file_uploader(
        "Upload Notes (PDF, DOCX, TXT, JPG, PNG)",
        type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
        key="summarizer_upload",
    )
    summary_type = st.selectbox("Summary Type", SUMMARY_TYPES, key="summarizer_type")
    if st.button("Summarize", key="summarize_btn"):
        if uploaded_file is None:
            st.error("Please upload a file first.")
        else:
            with st.spinner("🤖 Reading and summarizing..."):
                text = extract_text_from_any_upload(uploaded_file)
                if not text.strip():
                    st.error("Could not extract any readable text from this file.")
                else:
                    content, error = ai_summarize_notes(text, summary_type=summary_type)
                    if error:
                        st.error(error)
                    else:
                        st.markdown(content)


def render_ai_important_questions_tool():
    """AI Study Assistant > AI Important Questions."""
    st.markdown("### ⭐ AI Important Questions")
    uploaded_file = st.file_uploader(
        "Upload Notes (PDF, DOCX, TXT, JPG, PNG)",
        type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
        key="impq_upload",
    )
    if st.button("Generate Important Questions", key="impq_btn"):
        if uploaded_file is None:
            st.error("Please upload a file first.")
        else:
            with st.spinner("🤖 Reading notes and finding important questions..."):
                text = extract_text_from_any_upload(uploaded_file)
                if not text.strip():
                    st.error("Could not extract any readable text from this file.")
                else:
                    content, error = ai_generate_important_questions(text)
                    if error:
                        st.error(error)
                    else:
                        st.markdown(content)


def render_ai_quiz_generator_tool():
    """AI Study Assistant > AI Quiz Generator (interactive, scores itself)."""
    st.markdown("### 🧠 AI Quiz Generator")
    with st.form("ai_quiz_form"):
        topic = st.text_input("Subject / Topic", key="quiz_topic")
        num_q = st.number_input("Number of Questions", 1, 15, 5, key="quiz_num")
        difficulty = st.selectbox("Difficulty", DIFFICULTY_LEVELS, key="quiz_difficulty")
        submitted = st.form_submit_button("Generate Quiz")

    if submitted:
        if not topic:
            st.error("Please enter a subject/topic.")
        else:
            with st.spinner("🤖 Generating quiz..."):
                quiz, error = ai_generate_quiz(topic, int(num_q), difficulty)
            if error or not quiz:
                st.error(error or "Could not generate a valid quiz. Please try again.")
            else:
                st.session_state.quiz_data = quiz
                st.session_state.quiz_answers = {}
                st.session_state.quiz_submitted = False

    if st.session_state.quiz_data:
        st.markdown("---")
        st.markdown("#### Attempt the Quiz")
        for i, q in enumerate(st.session_state.quiz_data):
            st.markdown(f"**Q{i + 1}. {q['question']}**")
            selected = st.radio(
                "Choose one:",
                q["options"],
                key=f"quiz_q_{i}",
                label_visibility="collapsed",
            )
            st.session_state.quiz_answers[i] = selected

        if st.button("Submit Quiz", key="quiz_submit_btn"):
            score = 0
            for i, q in enumerate(st.session_state.quiz_data):
                correct_option = q["options"][q["correct_index"]]
                if st.session_state.quiz_answers.get(i) == correct_option:
                    score += 1
            st.session_state.quiz_submitted = True
            st.session_state.quiz_score = score

        if st.session_state.quiz_submitted:
            total = len(st.session_state.quiz_data)
            st.success(f"🎯 Your Score: {st.session_state.quiz_score} / {total}")
            for i, q in enumerate(st.session_state.quiz_data):
                correct_option = q["options"][q["correct_index"]]
                st.markdown(f"Q{i + 1}: Correct answer — **{correct_option}**")


def render_ai_flashcards_tool():
    """AI Study Assistant > AI Flash Cards (flip-through revision cards)."""
    st.markdown("### 🗂️ AI Flash Cards")
    with st.form("ai_flash_form"):
        topic = st.text_input("Subject / Topic", key="flash_topic")
        num_cards = st.number_input("Number of Flash Cards", 1, 20, 8, key="flash_num")
        submitted = st.form_submit_button("Generate Flash Cards")

    if submitted:
        if not topic:
            st.error("Please enter a subject/topic.")
        else:
            with st.spinner("🤖 Generating flash cards..."):
                cards, error = ai_generate_flashcards(topic, int(num_cards))
            if error or not cards:
                st.error(error or "Could not generate flash cards. Please try again.")
            else:
                st.session_state.flashcards_data = cards
                st.session_state.flashcard_index = 0
                st.session_state.flashcard_flipped = False

    if st.session_state.flashcards_data:
        st.markdown("---")
        cards = st.session_state.flashcards_data
        idx = st.session_state.flashcard_index
        card = cards[idx]
        st.markdown(f"**Card {idx + 1} of {len(cards)}**")

        if not st.session_state.flashcard_flipped:
            st.markdown(
                f'<div class="feature-card"><h4>❓ {card["question"]}</h4></div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="feature-card"><h4>✅ {card["answer"]}</h4></div>',
                unsafe_allow_html=True,
            )

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("⬅ Previous", key="flash_prev", use_container_width=True) and idx > 0:
                st.session_state.flashcard_index -= 1
                st.session_state.flashcard_flipped = False
                st.rerun()
        with c2:
            if st.button("🔄 Flip", key="flash_flip", use_container_width=True):
                st.session_state.flashcard_flipped = not st.session_state.flashcard_flipped
                st.rerun()
        with c3:
            if st.button("Next ➡", key="flash_next", use_container_width=True) and idx < len(cards) - 1:
                st.session_state.flashcard_index += 1
                st.session_state.flashcard_flipped = False
                st.rerun()


def render_ai_explain_topic_tool():
    """AI Study Assistant > AI Explain Topic."""
    st.markdown("### 💡 AI Explain Topic")
    topic = st.text_input("Enter a topic name", key="explain_topic_input")
    if st.button("Explain", key="explain_btn"):
        if not topic:
            st.error("Please enter a topic.")
        else:
            with st.spinner("🤖 Explaining..."):
                content, error = ai_explain_topic(topic)
            if error:
                st.error(error)
            else:
                st.markdown(content)


def render_ai_ask_questions_tool():
    """AI Study Assistant > AI Ask Questions.
    Upgraded (RAG): now supports two modes -
      1. 'Upload a specific file' - original behaviour, answers only
         from that one uploaded document.
      2. 'Search the whole repository (RAG)' - retrieves the most
         relevant chunks across every uploaded document using
         semantic_search_files()/build_rag_context(), answers using
         only that retrieved material, and shows the supporting
         sources + relevance scores underneath the answer
         (Explainable AI)."""
    st.markdown("### 💬 AI Ask Questions")
    mode = st.radio(
        "Where should I look for the answer?",
        ["Upload a specific file", "Search the whole repository (RAG)"],
        key="ask_mode",
    )

    if mode == "Upload a specific file":
        uploaded_file = st.file_uploader(
            "Upload Notes (PDF, DOCX, TXT, JPG, PNG)",
            type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
            key="ask_upload",
        )
        question = st.text_input("Your Question (e.g. 'What is NLP?')", key="ask_question_input")
        if st.button("Ask", key="ask_btn"):
            if uploaded_file is None or not question:
                st.error("Please upload notes and enter a question.")
            else:
                with st.spinner("🤖 Reading your notes..."):
                    text = extract_text_from_any_upload(uploaded_file)
                    if not text.strip():
                        st.error("Could not extract any readable text from this file.")
                    else:
                        content, error = ai_ask_from_notes(text, question)
                        if error:
                            st.error(error)
                        else:
                            st.markdown(content)
                        with st.expander("🔍 Debug: preview of extracted text sent to the AI"):
                            st.caption(f"Total extracted length: {len(text)} characters (first 12,000 are used).")
                            st.text(text[:2000])
    else:
        question = st.text_input("Your Question", key="ask_rag_question_input")
        if st.button("Ask Repository", key="ask_rag_btn"):
            if not question:
                st.error("Please enter a question.")
            else:
                with st.spinner("🤖 Searching the repository and thinking..."):
                    chunks = build_rag_context(question, top_k=5)
                    content, error = ai_rag_answer(question, chunks)
                if error:
                    st.error(error)
                else:
                    st.markdown(content)
                    if chunks:
                        with st.expander("📎 Sources used (Explainable AI)"):
                            for i, c in enumerate(chunks, start=1):
                                score_txt = f" — relevance {c['score']:.0%}" if c.get("score") is not None else ""
                                st.markdown(f"**Source {i}: {c['source']}{score_txt}**")
                                st.caption(c["text"][:300] + ("..." if len(c["text"]) > 300 else ""))


def render_ai_study_planner_tool():
    """AI Study Assistant > AI Study Planner."""
    st.markdown("### 🗓️ AI Study Planner")
    with st.form("ai_planner_form"):
        subject = st.text_input("Subject", key="planner_subject")
        exam_date = st.date_input("Exam Date", key="planner_exam_date")
        hours_per_day = st.number_input("Available Study Hours per day", 1, 12, 2, key="planner_hours")
        submitted = st.form_submit_button("Generate Study Plan")

    if submitted:
        if not subject:
            st.error("Please enter a subject.")
        elif exam_date <= datetime.now().date():
            st.error("Please choose a future exam date.")
        else:
            with st.spinner("🤖 Building your study timetable..."):
                plan, error = ai_generate_study_planner(subject, str(exam_date), int(hours_per_day))
            if error or not plan:
                st.error(error or "Could not generate a study plan. Please try again.")
            else:
                st.markdown("---")
                for day in plan:
                    st.markdown(
                        f"**{day.get('date', '')} (Day {day.get('day_number', '')})** — "
                        f"{day.get('hours', '')} hrs — {day.get('topics', '')}"
                    )


def render_ai_keyword_extraction_tool():
    """AI Study Assistant > AI Keyword Extraction."""
    st.markdown("### 🔑 AI Keyword Extraction")
    uploaded_file = st.file_uploader(
        "Upload Notes (PDF, DOCX, TXT, JPG, PNG)",
        type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
        key="keyword_upload",
    )
    if st.button("Extract Keywords", key="keyword_btn"):
        if uploaded_file is None:
            st.error("Please upload a file first.")
        else:
            with st.spinner("🤖 Extracting keywords..."):
                text = extract_text_from_any_upload(uploaded_file)
                if not text.strip():
                    st.error("Could not extract any readable text from this file.")
                else:
                    result, error = ai_extract_keywords(text)
                    if error or not result:
                        st.error(error or "Could not extract keywords. Please try again.")
                    else:
                        st.markdown("**Important Keywords:** " + ", ".join(result.get("keywords", [])))
                        st.markdown("**Technical Terms:** " + ", ".join(result.get("technical_terms", [])))
                        if result.get("definitions"):
                            st.markdown("**Definitions:**")
                            for d in result["definitions"]:
                                st.markdown(f"- **{d.get('term', '')}**: {d.get('definition', '')}")


def render_ai_knowledge_graph_tool():
    """AI Study Assistant > Knowledge Graph Generator (new)."""
    st.markdown("### 🧠 Knowledge Graph Generator")
    st.caption(
        "Generates a concept graph showing how sub-topics and related "
        "concepts connect for a subject/topic."
    )
    topic = st.text_input("Subject or Topic", key="kg_topic")
    if st.button("Generate Knowledge Graph", key="kg_btn"):
        if not topic:
            st.error("Please enter a subject or topic.")
        else:
            with st.spinner("🤖 Mapping out the concept graph..."):
                graph_data, error = ai_generate_knowledge_graph_data(topic)
            if error or not graph_data:
                st.error(error or "Could not generate a knowledge graph. Please try again.")
            else:
                render_knowledge_graph(graph_data)


def render_ai_resource_recommendation_tool():
    """AI Study Assistant > AI Resource Recommendations (new)."""
    st.markdown("### 🎯 AI Resource Recommendations")
    st.caption(
        "Finds related material already in the MBU Study Vault repository, "
        "plus AI-suggested sub-topics and search directions for further "
        "reading."
    )
    topic = st.text_input("Topic / Search Query", key="reco_topic")
    if st.button("Get Recommendations", key="reco_btn"):
        if not topic:
            st.error("Please enter a topic or search query.")
        else:
            with st.spinner("🤖 Finding related resources..."):
                repo_matches = semantic_search_files(topic, top_k=8)
                if not repo_matches:
                    repo_matches = search_files(topic)[:8]

            st.markdown("#### 📚 Related Resources in the Repository")
            if repo_matches:
                render_file_list(repo_matches)
            else:
                st.info("No related resources found in the repository yet.")

            with st.spinner("🤖 Thinking of related study directions..."):
                content, error = ai_recommend_search_directions(topic, repo_matches)
            st.markdown("#### 🧭 Suggested Study Directions")
            if error:
                st.error(error)
            else:
                st.markdown(content)


AI_STUDY_TOOLS = {
    "❓ AI Question Generator": render_ai_question_generator_tool,
    "📄 AI Question Paper Generator": render_ai_question_paper_generator_tool,
    "📝 AI Notes Summarizer": render_ai_notes_summarizer_tool,
    "⭐ AI Important Questions": render_ai_important_questions_tool,
    "🧠 AI Quiz Generator": render_ai_quiz_generator_tool,
    "🗂️ AI Flash Cards": render_ai_flashcards_tool,
    "💡 AI Explain Topic": render_ai_explain_topic_tool,
    "💬 AI Ask Questions": render_ai_ask_questions_tool,
    "🗓️ AI Study Planner": render_ai_study_planner_tool,
    "🔑 AI Keyword Extraction": render_ai_keyword_extraction_tool,
    "🧠 Knowledge Graph Generator": render_ai_knowledge_graph_tool,
    "🎯 AI Resource Recommendations": render_ai_resource_recommendation_tool,
}


def render_ai_study_assistant():
    """
    Hub page for the AI Study Assistant. Available to EVERY logged-in
    user - student or faculty - via a dropdown that switches between
    the AI tools, keeping each tool's form state isolated via unique
    widget keys.
    """
    st.markdown("## 🤖 AI Study Assistant")

    if not (st.session_state.logged_in or st.session_state.faculty_logged_in):
        st.warning("Please login to use the AI Study Assistant.")
        return

    st.caption("Available to every logged-in student and faculty member.")
    choice = st.selectbox("Choose a tool", list(AI_STUDY_TOOLS.keys()), key="ai_tool_choice")
    st.markdown("---")
    AI_STUDY_TOOLS[choice]()


# ======================================================================
# 17B. NEW PAGE - AI CHAT COMPANION (dedicated support chatbot)
# ======================================================================
def render_ai_chat_companion():
    st.markdown("## 💬 AI Study Companion")

    if not (st.session_state.logged_in or st.session_state.faculty_logged_in):
        st.warning("Please login to chat with the Study Companion.")
        return

    st.caption(
        "Your academic + motivational study companion. This is not a "
        "substitute for a counsellor, doctor, or trusted person - for "
        "anything serious, please also reach out to them."
    )

    st.markdown("**Quick actions**")
    qcols = st.columns(4)
    quick_actions = [
        ("Ask a Doubt", "I have a doubt, can you help me understand it? "),
        ("Explain a Topic", "Please explain this topic simply: "),
        ("Make a Study Plan", "Help me make a study plan for: "),
        ("Prepare for Exam", "Help me prepare for my exam on: "),
    ]
    qcols2 = st.columns(3)
    quick_actions2 = [
        ("Motivate Me", "I'm feeling low about my studies right now."),
        ("Quiz Me", "Quiz me with a few questions on: "),
        ("Ask From My Documents", "__RAG__"),
    ]

    prefill = None
    for col, (label, prompt_prefix) in zip(qcols, quick_actions):
        with col:
            if st.button(label, key=f"qa_{label}", use_container_width=True):
                prefill = prompt_prefix
    for col, (label, prompt_prefix) in zip(qcols2, quick_actions2):
        with col:
            if st.button(label, key=f"qa_{label}", use_container_width=True):
                prefill = prompt_prefix

    if prefill == "__RAG__":
        st.session_state["_companion_rag_mode"] = True
        st.info("Type your question below - it will be answered using the repository's uploaded documents.")
    elif prefill:
        st.session_state["_companion_rag_mode"] = False
        st.info(f"Quick action selected: continue typing your message below, e.g. \"{prefill}...\"")

    st.markdown("---")
    for msg in st.session_state.companion_chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    user_input = st.chat_input("Type your message...")

    if user_input:
        st.session_state.companion_chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        rag_context_text = None
        if st.session_state.get("_companion_rag_mode"):
            chunks = build_rag_context(user_input, top_k=4)
            if chunks:
                rag_context_text = "\n\n".join(
                    f"[{c['source']}]\n{c['text'][:1200]}" for c in chunks
                )

        with st.chat_message("assistant"):
            with st.spinner("🤖 Thinking..."):
                reply, error = ai_companion_reply(
                    st.session_state.companion_chat_history, rag_context=rag_context_text
                )
            if error:
                st.error(error)
            else:
                st.markdown(reply)
                st.session_state.companion_chat_history.append({"role": "assistant", "content": reply})

    if st.button("🧹 Clear Conversation"):
        st.session_state.companion_chat_history = []
        st.session_state["_companion_rag_mode"] = False
        st.rerun()


# ======================================================================
# 18. MAIN ROUTING
# ======================================================================
inject_css()
sidebar_nav()

current_page = st.session_state.page

if current_page == "Home":
    render_home()
elif current_page == "Signup":
    render_signup()
elif current_page == "Login":
    render_login()
elif current_page == "Dashboard":
    render_dashboard()
elif current_page == "Upload":
    render_upload()
elif current_page == "Search":
    render_search()
elif current_page == "Faculty Login":
    render_faculty_login()
elif current_page == "Faculty Signup":
    render_faculty_signup()
elif current_page == "Faculty Upload":
    render_faculty_upload()
elif current_page == "Faculty Dashboard":
    render_faculty_dashboard()
elif current_page == "AI Study Assistant":
    render_ai_study_assistant()
elif current_page == "AI Chat Companion":
    render_ai_chat_companion()
elif current_page == "Question Paper History":
    render_question_paper_history()
else:
    render_home()

st.markdown(
    """
    <div class="footer">
        Made with ❤️ for Mohan Babu University Students | MBU Study Vault © 2026
    </div>
    """,
    unsafe_allow_html=True,
)
